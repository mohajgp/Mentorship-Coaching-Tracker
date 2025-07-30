import streamlit as st
import pandas as pd
import plotly.express as px
from datetime import datetime
import pyperclip
from io import BytesIO
from docx import Document

# -------------------- PAGE CONFIG --------------------
st.set_page_config(
    page_title="KNCCI Mentorship Dashboard",
    layout="wide"
)

st.title("KNCCI Jiinue Mentorship Dashboard")
st.caption("Tracking Mentorship Sessions by Field Officers")

# -------------------- SETTINGS --------------------
OLD_FORM_URL = "https://docs.google.com/spreadsheets/d/107tWhbwBgF89VGCRnNSL4-_WrCIa68NGSPA4GkwVjDE/export?format=csv"
NEW_FORM_URL = "https://docs.google.com/spreadsheets/d/1CA7WvTkEUfeMyLuxhP91BgSWbkhV_S8V94kACj5LUMM/export?format=csv"

# -------------------- LOAD AND MERGE --------------------
@st.cache_data(ttl=300)
def load_raw_data():
    df_old = pd.read_csv(OLD_FORM_URL)
    df_new = pd.read_csv(NEW_FORM_URL)

    df_old.columns = df_old.columns.str.strip()
    df_new.columns = df_new.columns.str.strip()

    df_old['Form Version'] = 'Original'
    df_new['Form Version'] = 'New'

    df_old.rename(columns={
        'Timestamp': 'Timestamp',
        'County': 'County',
        'Gender': 'Gender',
        'Age': 'Age',
        'Business Phone Number': 'Phone Number'
    }, inplace=True)

    df_new.rename(columns={
        'Timestamp': 'Timestamp',
        '14. County of Business Location': 'County',
        '12. Gender of mentee (participant)': 'Gender',
        '11. Age of mentee (full years)': 'Age',
        '10. Mobile phone Number (Format: 2547XXXXXXXX)': 'Phone Number',
        '8. National ID (5 to 11 digits)': 'ID'
    }, inplace=True)

    for df in [df_old, df_new]:
        df['Timestamp'] = pd.to_datetime(df['Timestamp'], errors='coerce')
        df['County'] = df['County'].astype(str).str.strip().str.title()
        df['Gender'] = df['Gender'].astype(str).str.strip().str.title()
        df['Age'] = pd.to_numeric(df['Age'], errors='coerce')
        df['Phone Number'] = df['Phone Number'].astype(str).str.extract(r'(\d{9,12})')
        if 'ID' not in df.columns:
            df['ID'] = pd.NA

    for col in df_new.columns:
        if col not in df_old.columns:
            df_old[col] = pd.NA
    df_old = df_old[df_new.columns]

    merged_df = pd.concat([df_old, df_new], ignore_index=True)
    return merged_df, merged_df.shape[0]

# -------------------- ALL COUNTIES --------------------
all_counties_47 = [
    "Mombasa", "Kwale", "Kilifi", "Tana River", "Lamu", "Taita Taveta",
    "Garissa", "Wajir", "Mandera", "Marsabit", "Isiolo", "Meru", "Tharaka Nithi",
    "Embu", "Kitui", "Machakos", "Makueni", "Nyandarua", "Nyeri", "Kirinyaga",
    "Murang'a", "Kiambu", "Turkana", "West Pokot", "Samburu", "Trans Nzoia",
    "Uasin Gishu", "Elgeyo Marakwet", "Nandi", "Baringo", "Laikipia", "Nakuru",
    "Narok", "Kajiado", "Kericho", "Bomet", "Kakamega", "Vihiga", "Bungoma",
    "Busia", "Siaya", "Kisumu", "Homa Bay", "Migori", "Kisii", "Nyamira", "Nairobi"
]

# -------------------- LOAD --------------------
df_raw, total_raw_rows = load_raw_data()
if df_raw.empty:
    st.error("âŒ No data available! Please check both spreadsheets.")
    st.stop()

# -------------------- UNFILTERED COUNTY TOTALS --------------------
st.subheader("ğŸ“ Unfiltered Submissions by County (All Data)")
unfiltered_county_counts = df_raw.groupby('County').size().reset_index(name='Submissions (Unfiltered)')
fig_unfiltered_bar = px.bar(
    unfiltered_county_counts,
    x='County', y='Submissions (Unfiltered)', color='County',
    title='Unfiltered Number of Submissions by County'
)
st.plotly_chart(fig_unfiltered_bar, use_container_width=True)

st.subheader("ğŸ“Š Unfiltered County Submissions Data")
st.dataframe(unfiltered_county_counts)
unfiltered_csv = unfiltered_county_counts.to_csv(index=False).encode('utf-8')
st.download_button(
    label="â¬‡ï¸ Download Unfiltered County Submissions CSV",
    data=unfiltered_csv,
    file_name=f"Unfiltered_County_Submissions_{datetime.now().date()}.csv",
    mime='text/csv'
)

# -------------------- FILTERS --------------------
st.sidebar.header("ğŸ—“ï¸ Filter Sessions")
min_date = df_raw['Timestamp'].min().date()
max_date = df_raw['Timestamp'].max().date()

st.sidebar.markdown(f"ğŸ—“ï¸ **Earliest Submission**: {min_date}")
st.sidebar.markdown(f"ğŸ—“ï¸ **Latest Submission**: {max_date}")

date_range = st.sidebar.date_input("Select Date Range:", value=(min_date, max_date), min_value=min_date, max_value=max_date)
start_date, end_date = date_range if isinstance(date_range, tuple) else (date_range, date_range)

form_versions = df_raw['Form Version'].unique().tolist()
selected_versions = st.sidebar.multiselect("Select Form Version:", options=form_versions, default=form_versions)

counties = df_raw['County'].dropna().unique()
selected_counties = st.sidebar.multiselect("Select Counties:", options=sorted(counties), default=sorted(counties))

genders = df_raw['Gender'].dropna().unique()
selected_genders = st.sidebar.multiselect("Select Gender:", options=sorted(genders), default=sorted(genders))

filtered_df = df_raw[
    (df_raw['Timestamp'].dt.date >= start_date) &
    (df_raw['Timestamp'].dt.date <= end_date) &
    (df_raw['Form Version'].isin(selected_versions)) &
    (df_raw['County'].isin(selected_counties)) &
    (df_raw['Gender'].isin(selected_genders))
]

deduped_df = filtered_df.drop_duplicates(subset=['Phone Number', 'ID'])
total_unique_rows = deduped_df.shape[0]

# -------------------- METRICS --------------------
st.subheader("ğŸ“ˆ Summary Metrics")
col1, col2, col3, col4 = st.columns(4)
col1.metric("ğŸ“„ Filtered Raw Records", f"{filtered_df.shape[0]:,}")
col2.metric("âœ… Unique Records (Post-Dedup)", f"{total_unique_rows:,}")
col3.metric("ğŸ“Š Filtered Sessions", f"{total_unique_rows:,}")
col4.metric("ğŸ“ Counties Covered", deduped_df['County'].nunique())

# -------------------- YOUTH & FEMALE YOUTH ANALYSIS --------------------
st.subheader("ğŸ§‘â€ğŸ“ Youth Participation Overview (Ages 18â€“35)")

youth_df = deduped_df[(deduped_df['Age'] >= 18) & (deduped_df['Age'] <= 35)]
total_youths = youth_df.shape[0]
female_youths = youth_df[youth_df['Gender'].str.lower() == 'female'].shape[0]

percent_youth = (total_youths / total_unique_rows) * 100 if total_unique_rows else 0
percent_female_youth = (female_youths / total_unique_rows) * 100 if total_unique_rows else 0

coly1, coly2 = st.columns(2)
coly1.metric("ğŸ§‘ Youths (18â€“35)", f"{percent_youth:.2f}% ({total_youths:,})")
coly2.metric("ğŸ‘© Female Youths (18â€“35)", f"{percent_female_youth:.2f}% ({female_youths:,})")

# -------------------- AUTO SUMMARY --------------------
st.subheader("ğŸ“ Auto-Generated Summary Report")
no_submission_counties = [c for c in all_counties_47 if c not in deduped_df['County'].unique()]
summary_text = f"""ğŸ“… **Date Range**: {start_date} to {end_date}
ğŸ“„ **Filtered Raw Records**: {filtered_df.shape[0]:,}
âœ… **Unique Records (Post-Dedup)**: {total_unique_rows:,}
ğŸ“Š **Filtered Sessions**: {total_unique_rows:,}
ğŸ“ **Counties Covered**: {deduped_df['County'].nunique()}
ğŸ§‘ **Youths (18â€“35)**: {total_youths:,} ({percent_youth:.2f}%)
ğŸ‘© **Female Youths (18â€“35)**: {female_youths:,} ({percent_female_youth:.2f}%)
ğŸš« **Counties with No Submissions**: {len(no_submission_counties)} ({', '.join(no_submission_counties)})"""

st.text_area("ğŸ“‹ Copy this Summary for Emailing:", value=summary_text, height=240)
if st.button("ğŸ“‹ Copy to Clipboard"):
    pyperclip.copy(summary_text)
    st.success("âœ… Summary copied to clipboard!")

# -------------------- EXPORT TO WORD --------------------
st.subheader("ğŸ“„ Export Summary to Word Document")
if st.button("â¬‡ï¸ Generate Word Report"):
    doc = Document()
    doc.add_heading("KNCCI Jiinue Mentorship Summary Report", level=1)
    doc.add_paragraph(summary_text)
    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    st.download_button(
        label="ğŸ“¥ Download Word Report",
        data=buffer,
        file_name=f"Mentorship_Summary_{datetime.now().date()}.docx",
        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    )

# -------------------- COUNTY CHART + TABLE --------------------
st.subheader("ğŸ“ Filtered Submissions by County")
county_counts = deduped_df.groupby('County').size().reset_index(name='Submissions')
fig_bar = px.bar(county_counts, x='County', y='Submissions', color='County', title='Number of Submissions by County')
st.plotly_chart(fig_bar, use_container_width=True)

st.subheader("ğŸ“Š Filtered County Submissions Data")
st.dataframe(county_counts)
csv_data = county_counts.to_csv(index=False).encode('utf-8')
st.download_button(
    label="â¬‡ï¸ Download Filtered County Submissions CSV",
    data=csv_data,
    file_name=f"County_Submissions_{datetime.now().date()}.csv",
    mime='text/csv'
)

# -------------------- DAILY TREND --------------------
st.subheader("ğŸ“† Submissions Over Time")
daily_counts = deduped_df.groupby(deduped_df['Timestamp'].dt.date).size().reset_index(name='Submissions')
fig_time = px.line(daily_counts, x='Timestamp', y='Submissions', title='Daily Submission Trend')
st.plotly_chart(fig_time, use_container_width=True)

# -------------------- NON-SUBMITTING COUNTIES --------------------
st.subheader("ğŸš« Counties with No Submissions")
if no_submission_counties:
    st.error(f"ğŸš« Counties with **NO** Submissions: {', '.join(no_submission_counties)}")
else:
    st.success("âœ… All counties have submissions in selected date range.")

# -------------------- DATA TABLES --------------------
st.subheader("ğŸ“„ Filtered Raw Records (With Duplicates)")
st.dataframe(filtered_df)

st.subheader("âœ… Cleaned Unique Records (Post-Filter)")
st.dataframe(deduped_df)

# -------------------- EXPORT MERGED DATA --------------------
st.subheader("â• Merged Full Data")
full_csv = df_raw.to_csv(index=False).encode('utf-8')
st.download_button(
    label="ğŸ“¥ Download Merged CSV",
    data=full_csv,
    file_name=f"Mentorship_Merged_Data_{datetime.now().date()}.csv",
    mime='text/csv'
)
